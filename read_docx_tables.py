import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    
    # Read paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Read tables
    for table in doc.tables:
        full_text.append("--- TABLE START ---")
        for row in table.rows:
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            full_text.append(" | ".join(row_data))
        full_text.append("--- TABLE END ---")
        
    return '\n'.join(full_text)

if __name__ == "__main__":
    text = read_docx(sys.argv[1])
    with open("protocolo_tables.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Done")
